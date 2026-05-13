from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Codex\Duals")
OUTPUT_DIR = ROOT / "docs"
OUTPUT_PATH = OUTPUT_DIR / "Gestio_de_les_practiques_Duals_Document_Final_Fase_Estudi.docx"
LOGO_PATH = ROOT / "Imatges" / "logo.png"

NAVY = RGBColor(30, 44, 69)
RED = RGBColor(214, 17, 33)
GRAY = RGBColor(85, 85, 85)
LIGHT = RGBColor(244, 246, 249)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_page_background(section):
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    run = p.add_run(text)
    if level == 1:
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.color.rgb = RED
    return p


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.size = Pt(10.5)


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = NAVY
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(10.5)

    hdr = table.rows[0].cells
    hdr[0].text = "Bloc"
    hdr[1].text = "Descripció"
    set_repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, "1E2C45")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        set_cell_shading(cells[0], "F4F6F9")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 2"].font.bold = True

    for section in doc.sections:
        set_page_background(section)

    cover = doc.sections[0]
    header = cover.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH.exists():
        header.add_run().add_picture(str(LOGO_PATH), width=Cm(4.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Proposta Funcional del Projecte Gestió de les pràctiques Duals")
    r.font.name = "Aptos Display"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Fase 1 i línies de futur")
    r.font.size = Pt(16)
    r.font.color.rgb = RED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Document de treball per a presentació, consens i priorització amb el Departament"
    )
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = GRAY

    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    info.autofit = False
    info.columns[0].width = Cm(4.0)
    info.columns[1].width = Cm(9.5)
    meta = [
        ("Projecte", "Gestió de les pràctiques Duals"),
        ("Ubicació prevista", r"D:\Codex\Duals"),
        ("Àmbit", "Gestió de pràctiques Duals dels cicles formatius"),
        ("Objectiu", "Definir l’abast funcional de la primera fase i les evolucions futures"),
    ]
    for row, values in zip(info.rows, meta):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]
        set_cell_shading(row.cells[0], "F4F6F9")
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_page_break()

    add_heading(doc, "1. Resum executiu", 1)
    add_paragraph(
        doc,
        "El projecte Gestió de les pràctiques Duals planteja la substitució progressiva de la gestió actual basada en fitxers Excel per una aplicació web corporativa que centralitzi la informació d’alumnes, professorat, empreses, convenis, seguiment i documentació associada a les pràctiques Duals.",
    )
    add_paragraph(
        doc,
        "La fase 1 té com a finalitat posar en marxa una primera versió operativa, usable per administració, equip directiu, tutors i alumnat, capaç de treballar per cursos acadèmics, suportar múltiples convenis per alumne i mantenir traçabilitat real del procés.",
    )

    add_heading(doc, "2. Objectius del sistema", 1)
    add_bullet_list(
        doc,
        [
            "Centralitzar la gestió de pràctiques Duals en una única plataforma.",
            "Eliminar dependència operativa de l’Excel com a eina principal de seguiment.",
            "Permetre governança per cursos acadèmics, cicles, classes i tutors.",
            "Facilitar la participació de l’alumnat en l’actualització de dades autoritzades.",
            "Disposar d’un historial de seguiment, documents i convenis traçable.",
            "Preparar una base sòlida per futures funcionalitats intel·ligents i d’automatització.",
        ],
    )

    add_heading(doc, "2.1 Principis de construcció", 2)
    add_bullet_list(
        doc,
        [
            "La solució s’ha de construir com una solució estàndard editable amb Visual Studio 2026 o versions posteriors.",
            "En la implementació s’ha de prioritzar la simplicitat del codi, la llegibilitat i la facilitat de manteniment.",
            "S’evitarà sobrearquitectura innecessària i es prioritzarà una estructura clara i estable per al futur equip de manteniment.",
        ],
    )

    add_heading(doc, "3. Problema actual i necessitat", 1)
    add_paragraph(
        doc,
        "La gestió actual presenta dependència de fulls Excel, poca escalabilitat, dificultat per mantenir historial estructurat, risc d’inconsistències entre anys acadèmics i una capacitat limitada per compartir informació de forma segura entre diferents perfils d’usuari.",
    )
    add_paragraph(
        doc,
        "A més, la necessitat de treballar amb alumnat, professorat i empreses, així com amb convenis i seguiment documental, fa recomanable un sistema amb accés web, permisos per rol, registre d’activitat i capacitat d’evolució futura.",
    )

    add_heading(doc, "4. Abast funcional de la fase 1", 1)
    add_two_col_table(
        doc,
        [
            ("Estructura acadèmica", "Gestió de cursos acadèmics, cicles, classes i assignació de tutors."),
            ("Alumnat", "Fitxa única d’alumne, matrícula anual per curs, CV i dades complementàries."),
            ("Professorat", "Fitxa pròpia, autenticació i gestió de classes assignades."),
            ("Empreses", "Llista compartida, contactes i tutors d’empresa reutilitzables."),
            ("Convenis", "Múltiples convenis per alumne i any, amb restricció d’un únic conveni obert."),
            ("Seguiment", "Historial d’esdeveniments, actes i agenda de reunions/trucades."),
            ("Documents", "CVs, fitxers de conveni i repositori comú de plantilles."),
            ("Dashboards", "Resums per professor, gestor i administrador."),
            ("Administració", "Configuració, enviament manual de credencials i consulta de logs."),
        ],
    )

    add_heading(doc, "5. Actors i rols", 1)
    add_paragraph(doc, "El sistema contemplarà quatre perfils principals d’ús.")
    add_bullet_list(
        doc,
        [
            "Administrador: visió i operativa global completa.",
            "Gestor: visió global equivalent a l’administrador, amb capacitat d’edició només dins el seu àmbit docent si també és tutor.",
            "Professor/Tutor: gestiona les seves classes i els alumnes assignats.",
            "Alumne: consulta i completa únicament les dades autoritzades pel seu tutor.",
        ],
    )

    add_heading(doc, "6. Model acadèmic i de dades", 1)
    add_paragraph(
        doc,
        "L’alumne tindrà una única fitxa mestra persistent al llarg del temps. La seva situació acadèmica s’expressarà per mitjà de registres anuals vinculats a curs acadèmic, classe i cicle. Això permet gestionar promocions de primer a segon, especialitats diferents i casos de repetidors sense duplicar la identitat de l’alumne.",
    )
    add_bullet_list(
        doc,
        [
            "Un professor pertany a un cicle.",
            "Un professor pot gestionar una o més classes del seu cicle.",
            "Un alumne pertany a una classe per curs acadèmic.",
            "Un alumne pot tenir més d’un conveni en el mateix curs acadèmic.",
            "Un alumne només pot tenir un conveni obert per curs acadèmic.",
        ],
    )

    add_heading(doc, "7. Gestió de l’alumnat", 1)
    add_paragraph(
        doc,
        "L’alumne accedirà amb el correu electrònic de l’escola i una contrasenya generada pel sistema. L’enviament de credencials serà sempre una acció manual decidida per l’administrador, per a un, diversos o tots els alumnes seleccionats.",
    )
    add_paragraph(
        doc,
        "L’alumne només podrà modificar camps que no provinguin de l’Excel oficial del centre. El tutor serà qui obrirà i tancarà el període d’edició del formulari. L’alumne només podrà pujar el seu CV; la foto no serà editable des del seu perfil.",
    )

    add_heading(doc, "8. Gestió del professorat", 1)
    add_paragraph(
        doc,
        "El professor accedirà amb el correu electrònic de l’escola i una contrasenya generada automàticament. Aquesta contrasenya podrà ser modificada pel propi professor o per l’administrador. El professor disposarà de fitxa pròpia, incloent-hi fotografia carregada individualment.",
    )

    add_heading(doc, "9. Empreses i tutors d’empresa", 1)
    add_paragraph(
        doc,
        "La llista d’empreses serà comuna a tot el professorat. Una empresa podrà tenir diferents tutors d’empresa segons el conveni i l’alumne, i aquests tutors quedaran enregistrats per poder reutilitzar-se posteriorment en nous convenis.",
    )

    add_heading(doc, "10. Convenis i seguiment", 1)
    add_paragraph(
        doc,
        "Els convenis tindran estat Obert, Tancat o Finalitzat. Aquest estat s’ha de diferenciar del seguiment operatiu del procés, que reflectirà l’evolució real de cada alumne: enviament de CV, resposta de l’empresa, entrevista, selecció, conveni pendent, pràctiques en curs, finalització, etc.",
    )
    add_paragraph(
        doc,
        "La llista d’esdeveniments serà configurable. Per garantir l’evolució multilingüe del sistema, cada esdeveniment s’haurà de modelar amb una clau estable i textos localitzats per idioma, de manera que més endavant es pugui afegir l’anglès sense alterar ni trencar l’històric registrat.",
    )

    add_heading(doc, "11. Actes i agenda", 1)
    add_paragraph(
        doc,
        "El sistema haurà de permetre registrar actes de trucades, reunions o visites a empreses. Aquests registres hauran de poder ser planificats dins l’aplicació, de manera que es disposi d’una agenda bàsica amb seguiment d’acords, observacions i properes accions.",
    )

    add_heading(doc, "12. Documents i repositori compartit", 1)
    add_paragraph(
        doc,
        "La fase 1 inclourà la gestió de CVs, fitxers associats als convenis i altres documents del procés. També s’habilitarà un repositori compartit de fitxers i plantilles per a tot el professorat, sense versionat en aquesta fase.",
    )
    add_paragraph(
        doc,
        "Es considera també viable adjuntar correus electrònics com a evidències del procés, especialment en relació amb alumnat, convenis o seguiments. En una primera etapa es recomana permetre’n l’adjunció manual com a fitxer o evidència documental, deixant preparada l’arquitectura per a una futura integració més profunda amb Gmail.",
    )

    add_heading(doc, "13. Dashboards i informació resumida", 1)
    add_bullet_list(
        doc,
        [
            "Professor: resum dels alumnes de les seves classes, amb estat del procés, tipus i situació respecte a conveni.",
            "Gestor: la mateixa visió en mode global, amb possibilitat de lectura transversal.",
            "Administrador: visió completa, incloent configuració, imports i supervisió del sistema.",
        ],
    )

    add_heading(doc, "14. Logs i traçabilitat", 1)
    add_paragraph(
        doc,
        "L’administrador haurà de poder consultar els logs de la solució per bloc funcional i criticitat. Els blocs mínims recomanables són autenticació, imports, convenis, fitxers, correu, dashboards, configuració i sistema. Els nivells mínims seran Debug, Info, Warning, Error i Critical.",
    )

    add_heading(doc, "15. Comunicacions", 1)
    add_paragraph(
        doc,
        "L’enviament de correus no serà mai automàtic. Serà una acció manual llançada per l’administrador. La configuració del servidor de correu haurà de ser editable des del sistema i compatible, en una primera etapa, amb Gmail via SMTP.",
    )

    add_heading(doc, "16. Multidioma i PWA", 1)
    add_paragraph(
        doc,
        "La interfície es publicarà en català per defecte i castellà com a segon idioma, amb arquitectura preparada per incorporar anglès i altres idiomes. La web funcionarà com a PWA instal·lable, però exclusivament en mode online.",
    )

    add_heading(doc, "17. Punt d’entrada i desplegament pilot", 1)
    add_paragraph(
        doc,
        "El punt d’entrada de la web haurà de ser configurable. En una primera fase pilot es preveu la publicació a l’adreça https://tuteapps.ddns.net:4444, mantenint l’arquitectura preparada per futures variacions d’host, port o domini.",
    )

    add_heading(doc, "17.1 Persistència i criteri de portabilitat", 2)
    add_paragraph(
        doc,
        "La solució s’implantarà inicialment sobre PostgreSQL integrat dins del mateix docker compose. Aquesta decisió respon a criteris d’obertura tecnològica, bon encaix amb entorns Linux contenidoritzats i sostenibilitat a llarg termini.",
    )
    add_paragraph(
        doc,
        "Tot i això, la persistència s’haurà de dissenyar de manera que la base de dades es pugui moure sense complicacions a una instància externa de PostgreSQL. Addicionalment, es procurarà minimitzar dependències específiques del motor per facilitar, si mai cal, una futura adaptació controlada a altres motors compatibles amb EF Core.",
    )

    add_heading(doc, "18. Propostes clares de millora", 1)
    add_bullet_list(
        doc,
        [
            "Separar explícitament alumne, matrícula anual i conveni per evitar rigidesa futura.",
            "Distingir estat del conveni i estat del procés per no barrejar control documental i seguiment operatiu.",
            "Afegir estat del formulari d’alumne: no obert, obert, completat, revisat.",
            "Fer de la promoció de curs un procés guiat i no una operació manual dispersa.",
            "Registrar auditoria funcional d’accions sensibles, a més dels logs tècnics.",
            "Convertir esdeveniments i altres catàlegs en dades configurables des del principi.",
        ],
    )

    add_heading(doc, "19. Línies futures recomanades", 1)
    add_bullet_list(
        doc,
        [
            "Integració d’un assistent de consulta en llenguatge natural per explotar dades i documents.",
            "Cerca semàntica i resum de reunions, convenis i seguiments.",
            "Suggeriments intel·ligents sobre alumnat pendent, empreses sense resposta o convenis en risc.",
            "Automatització controlada de comunicacions i recordatoris.",
            "Portal d’empresa i fluxos documentals més avançats.",
        ],
    )
    add_paragraph(
        doc,
        "La integració futura d’un assistent de llenguatge natural es considera viable sempre que s’apliquin controls estrictes de permisos, traçabilitat i vincle amb la informació font. La recomanació és abordar-lo com a evolució posterior, un cop la dada estructural i documental estigui consolidada.",
    )

    add_heading(doc, "19.1 Incorporació d’MCP des de l’inici", 2)
    add_paragraph(
        doc,
        "Es proposa incorporar MCP des de la fase inicial com a capa d’extensió i interoperabilitat, però sense fer dependre el nucli funcional del sistema d’aquesta tecnologia. El sistema principal continuarà essent una solució .NET simple, editable des de Visual Studio, mentre que MCP actuarà com a via controlada per a consultes avançades, integració amb assistants i evolucions futures.",
    )
    add_bullet_list(
        doc,
        [
            "MCP s’orientarà inicialment a consultes de lectura i recuperació d’informació.",
            "La lògica funcional exposada per MCP haurà de reutilitzar els mateixos serveis interns que la web i l’API.",
            "S’hauran de respectar estrictament els permisos per rol i àmbit de dades.",
            "Es desaconsella començar amb operacions d’escriptura lliure sobre convenis o dades crítiques.",
        ],
    )
    add_paragraph(
        doc,
        "Aquesta decisió deixa el projecte preparat per a futures consultes en llenguatge natural, assistants interns, automatitzacions i integració d’eines externes, sense comprometre la simplicitat de l’arquitectura base.",
    )

    add_heading(doc, "20. Recomanació final", 1)
    add_paragraph(
        doc,
        "Abans d’iniciar la construcció, es recomana consensuar aquest document i tancar una versió curta i formal de l’abast de fase 1. Donat el nombre de matisos detectats, disposar d’un acord funcional explícit reduirà canvis de criteri i facilitarà una implementació molt més controlada.",
    )

    add_heading(doc, "21. Pauta definitiva d’implementació", 1)
    add_paragraph(
        doc,
        "Es proposa abordar la implementació en iteracions curtes, mantenint sempre un criteri de simplicitat, portabilitat i compatibilitat plena amb Visual Studio 2026 o posteriors.",
    )
    add_bullet_list(
        doc,
        [
            "Iteració 1: crear la solució base, estructura de projectes, configuració comuna, branding inicial i desplegament docker mínim amb PostgreSQL, Redis, web, API i Nginx.",
            "Iteració 2: implementar autenticació, rols, canvi de contrasenya i base de permisos per Administrador, Gestor, Professor/Tutor i Alumne.",
            "Iteració 3: model acadèmic complet amb cursos acadèmics, cicles, classes, professorat, alumnat i matrícules anuals.",
            "Iteració 4: empreses, tutors d’empresa reutilitzables i expedients de conveni amb restricció d’un únic conveni obert per alumne i any.",
            "Iteració 5: seguiment, esdeveniments configurables, actes i agenda de reunions/trucades.",
            "Iteració 6: documents, CVs, repositori compartit i adjunció manual de correus com a evidències del procés.",
            "Iteració 7: formulari d’alumne, obertura i tancament per tutor, i enviament manual de credencials i comunicacions.",
            "Iteració 8: imports d’Excel i ZIP, dashboards per rol, logs administratius i estabilització funcional.",
            "Iteració 9: activació de MCP en mode controlat per a consultes de lectura i explotació segura de la informació.",
            "Iteració 10: proves integrades, documentació tècnica, preparació de publicació i pilot a l’adreça prevista.",
        ],
    )
    add_paragraph(
        doc,
        "Cada iteració hauria de tancar-se amb validació funcional, revisió de permisos, prova de desplegament sobre docker compose i verificació explícita de mantenibilitat del codi.",
    )

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_background(section)
    add_heading(doc, "Annex. Decisions ja consensuades", 1)
    add_bullet_list(
        doc,
        [
            "Projecte ubicat a D:\\Codex\\Duals i publicat en un únic repositori GitHub.",
            "Branding basat en els recursos de la carpeta Imatges.",
            "Fitxers originals d’exemple centralitzats a la carpeta Exemples.",
            "Accés d’alumnes i professors amb correu de l’escola.",
            "Enviament manual de credencials per part de l’administrador.",
            "Foto d’alumne no editable per l’alumne.",
            "Foto de professor amb càrrega individual.",
            "Llista d’empreses comuna a tots els professors.",
            "Repositori comú de plantilles sense versionat.",
            "PWA només online.",
            "Arquitectura preparada per suportar MCP des del principi.",
            "Solució pensada per ser editable en Visual Studio 2026 o posteriors.",
            "Base de dades inicial PostgreSQL dins docker compose, preparada per externalitzar-se fàcilment.",
        ],
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
